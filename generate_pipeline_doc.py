import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_pipeline_docx():
    doc = Document()
    
    # Page setup (Letter / A4)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    PRIMARY = RGBColor(15, 23, 42)        # Midnight Slate (#0F172A)
    ACCENT_BLUE = RGBColor(14, 116, 144)  # Cyan/Teal (#0E7490)
    ACCENT_PURPLE = RGBColor(99, 102, 241)# Indigo (#6366F1)
    DARK_TEXT = RGBColor(30, 41, 59)
    MUTED_TEXT = RGBColor(100, 116, 139)
    GOLD_ACCENT = RGBColor(180, 83, 9)
    CODE_BG = "F1F5F9"
    CALLOUT_BG = "FEF3C7"
    DARK_HEADER_BG = "0F172A"

    # Document Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(2)
    r_main = title_p.add_run("ORBIT-OASIS: TECHNICAL PIPELINE ENCYCLOPEDIA")
    r_main.font.size = Pt(21)
    r_main.font.bold = True
    r_main.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    r_sub = sub_p.add_run("Deep-Dive Algorithmic Specifications, Tensor Transforms, Mathematical Proofs & Defense Scripts")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = ACCENT_BLUE

    def add_section_header(title, route, agent):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        cell = table.cell(0, 0)
        set_cell_background(cell, DARK_HEADER_BG)
        set_cell_margins(cell, top=130, bottom=130, left=160, right=160)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(title)
        r_title.font.size = Pt(13)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(255, 255, 255)

        p_sub = cell.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(2)
        p_sub.paragraph_format.space_after = Pt(0)
        r_s = p_sub.add_run(f"🌐 UI Route: {route}   |   🤖 Authority: {agent}")
        r_s.font.size = Pt(9)
        r_s.font.bold = True
        r_s.font.color.rgb = RGBColor(147, 197, 253)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_subheading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = ACCENT_PURPLE
        return p

    def add_body(text, space_after=4, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK_TEXT
        r.font.bold = bold
        r.font.italic = italic
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        rb = p.add_run(bold_prefix)
        rb.font.size = Pt(9.5)
        rb.font.bold = True
        rb.font.color.rgb = DARK_TEXT
        rt = p.add_run(text)
        rt.font.size = Pt(9.5)
        rt.font.color.rgb = DARK_TEXT

    def add_code_pipeline(pipeline_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        cell = table.cell(0, 0)
        set_cell_background(cell, CODE_BG)
        set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(pipeline_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(8)
        r.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_script_box(script_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        cell = table.cell(0, 0)
        set_cell_background(cell, CALLOUT_BG)
        set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r_lbl = p.add_run("🎙️ Verbatim Technical Pitch Script: ")
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = GOLD_ACCENT
        r_txt = p.add_run(f'"{script_text}"')
        r_txt.font.size = Pt(9)
        r_txt.font.italic = True
        r_txt.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Master Table
    add_subheading("Master Quick-Reference Table")
    q_table = doc.add_table(rows=8, cols=3)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    q_table.autofit = False
    col_w = [Inches(1.8), Inches(1.8), Inches(3.4)]
    for i, w in enumerate(col_w):
        q_table.columns[i].width = w

    headers = ["When on this Page...", "Point at this UI Element...", "Explain this Behind-the-Scenes Tech..."]
    for i, h in enumerate(headers):
        cell = q_table.rows[0].cells[i]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    table_data = [
        ("3D Crime Scene (/crime-scene)", "3D room, Trajectory laser, Measure tool", "NLP parses text into 3D bounding coordinates, and Three.js raycaster computes inverse ballistic trajectory vectors."),
        ("Deepfake Detection (/deepfake)", "Ensemble AI score, Grad-CAM heatmap", "Dual-model neural ensemble; Grad-CAM backpropagates gradients to highlight sub-pixel synthetic blending seams."),
        ("Biometrics (/biometric)", "Glowing radial Score Ring, Suspect card", "Converts fingerprint minutiae into graph adjacency matrices and extracts Gabor wavelets from polar-unwrapped iris contours."),
        ("Smart Assignment (/assignment)", "Recommended officers list & match %", "5-factor algorithmic optimization matrix balancing skill match, distance in km, caseload capacity, and historic success."),
        ("Evidence (/evidence)", "Degradation status tags & half-life", "Thermodynamic biochemical decay equations calculating sample viability half-life based on storage conditions."),
        ("Audit Trail (/audit-trail)", "Timeline hashes & green verified badges", "SHA-256 Merkle chain-of-custody logging where every state transition is cryptographically sealed for court admissibility."),
        ("Reports & Collab (/reports)", "1-click legal PDF & tactical board", "Automated juridical dossier compiler pulling data across all modules with multi-signature cryptographic sign-off.")
    ]

    for row_idx, r_data in enumerate(table_data):
        row_cells = q_table.rows[row_idx + 1].cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(r_data):
            set_cell_background(row_cells[col_idx], bg_col)
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8)
            r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: 3D Crime Scene
    add_section_header("SECTION 1: 3D Volumetric Crime Scene & Ballistics Solver", "client/pages/CrimeScene.tsx", "Agent Apex-Vision (Spatial Engine)")
    add_code_pipeline("""[Natural Language Incident Prompt]
        │
        ▼
[Stage 1: Spatial Semantic Entity Tokenization] -> Regex Boundary & Orientation Extraction
        │
        ▼
[Stage 2: Deterministic Pseudo-Random Seed Hash] -> Uniform Spatial Jitter [-0.4, 0.4]
        │
        ▼
[Stage 3: 3D Procedural Mesh & Hierarchy Assembly] -> Three.js WebGL Volumetric Scene
        │
        ▼
[Stage 4: Inverse Ballistic Raycaster Trajectory] -> R(t) = P_orig + t*D_dir => Euclidean Distance & Laser""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("Room Bounding Hierarchy: ", "Standard configs: office (12m x 9m x 3.5m), warehouse (18m x 14m x 6m), bedroom (9m x 8m x 3.2m). Wall offsets defined along +/- W/2 and +/- D/2.")
    add_bullet("Seeded PRNG Formula: ", "h = (sum(c_i * 31^(n-1-i))) mod 10000; jitter = (h mod 10000)/10000 - 0.5. Prevents evidence clustering across generations.")
    add_bullet("Raycaster Trajectory Math: ", "Computes normalized vector D = (P2 - P1) / ||P2 - P1||; Euclidean distance d = sqrt(Δx² + Δy² + Δz²); Pitch θ = arcsin(Δy/d); Yaw φ = atan2(Δz, Δx).")
    add_script_box("Judges, in our 3D Crime Scene module, Agent Apex-Vision takes unstructured incident text and performs spatial entity extraction. It maps boundary directions to 3D coordinate bounding boxes in Three.js and WebGL. When calculating trajectories, our raycaster resolves the Euclidean distance tensor between bullet casings and points of impact, computing elevation and azimuth vectors to determine the exact conical origin of the shooter.")

    # Section 2: Deepfake
    add_section_header("SECTION 2: Deepfake Detection & Frequency-Domain Heatmaps", "client/pages/DeepfakeDetection.tsx (Port 8000)", "Agent Apex-Vision (Media Forensics)")
    add_code_pipeline("""[Uploaded Image / Video File]
        │
        ▼
[Stage 1: Frame Ingestion & Normalization] -> OpenCV decodes [B x 3 x 224 x 224] tensors
        │
        ▼
[Stage 2: Dual Neural Network Ensemble] -> S_ens = σ(w1*f_prim(X) + w2*f_sec(X))
        │
        ▼
[Stage 3: Grad-CAM Sub-Pixel Heatmap] -> α_k^c = (1/Z) * sum(∂y^c / ∂A^k) => ReLU localization
        │
        ▼
[Stage 4: Base64 RGBA Heatmap & Video Timeline] -> Fake Frame Ratio = (1/N) * sum(I(S_t > τ))""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("Dual-Model Ensemble: ", "Combines primary structural boundary model (w1=0.65) with secondary high-frequency texture model (w2=0.35).")
    add_bullet("Grad-CAM Backpropagation: ", "Computes gradients of prediction score with respect to convolutional feature map activations: L_GradCAM = ReLU(sum(α_k * A^k)). Isolates sub-pixel GAN grid lines and diffusion blending seams.")
    add_bullet("Temporal Aggregation: ", "Samples keyframes over clip duration, computing fake-frame and AI-frame percentage ratios across the temporal timeline.")
    add_script_box("Judges, our deepfake detection pipeline doesn't rely on standard binary classifiers. Agent Apex-Vision executes a dual-model neural ensemble. We extract convolutional activations from the final layer and compute Grad-CAM heatmaps by backpropagating the gradients of the synthetic class logit. This isolates sub-pixel diffusion artifacts and GAN up-sampling checkerboard noise, generating frame-by-frame temporal fake ratios across entire video clips.")

    # Section 3: Biometrics
    add_section_header("SECTION 3: Multi-Modal Biometrics (Iris & Fingerprints)", "client/pages/Biometric.tsx (Port 8002)", "Agent Bio-Topology (Biometric Transformer)")
    add_code_pipeline("""[Biometric Input: Latent Fingerprint OR High-Res Iris Image]
        │
        ├── Modality A (Iris): Daugman Polar Normalization -> 2D Complex Gabor Wavelet -> 2048-bit IrisCode -> Fractional Hamming Distance
        │
        └── Modality B (Fingerprint): Crossing Number CN(P) -> Non-Euclidean Minutiae Graph G=(V,E) -> Hungarian Bipartite Matching
        │
        ▼
[Radial Score Ring Engine & Suspect Record] -> Score = (1 - Metric) * 100 => SVG Ring & Criminal Record Card""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("Daugman Integro-Differential Operator: ", "Maximizes radial contour gradient: max |G_σ(r) * (∂/∂r) ∮ (I(x,y)/2πr) ds|. Unwraps circular iris into polar strip I(r, θ).")
    add_bullet("2D Gabor Wavelet IrisCode: ", "Quantizes real and imaginary filter outputs into a 2048-bit binary vector, matching via fractional Hamming Distance: HD = ||(CodeA ⊕ CodeB) ∩ Mask|| / ||Mask||.")
    add_bullet("Crossing Number Minutiae Graphs: ", "CN(P) = 0.5 * sum(|Pi - Pi+1|). CN=1 for ridge termination, CN=3 for bifurcation. Builds Delaunay graphs matched via Hungarian bipartite optimization (Kuhn-Munkres).")
    add_script_box("Judges, on our Biometric Analysis tab, Agent Bio-Topology runs a dual-modality neural pipeline. For iris scans, we perform Daugman rubber-sheet normalization and extract phase coefficients via 2D Gabor wavelets into a 2048-bit IrisCode compared via fractional Hamming distances. For fingerprints, we extract Crossing Number ridge bifurcations, constructing non-Euclidean topological graphs matched via Hungarian bipartite optimization, achieving rotation-invariant matching in under 45 milliseconds.")

    # Section 4: Smart Assignment
    add_section_header("SECTION 4: Smart Case Assignment & Predictive Resolution", "client/pages/Assignment.tsx & Cases.tsx (Port 8001)", "Agent Nexus-Decision (Predictive Allocator)")
    add_code_pipeline("""[Incoming Case: Crime Type, Priority, Evidence Types, GPS Location]
        │
        ▼
[Stage 1: Multi-Variable Case Complexity Estimator] -> Days = Base_Duration * C_mult * (1 + 0.15*Num_Evidence)
        │
        ▼
[Stage 2: Haversine Geodesic Distance Matrix] -> d = 2R * atan2(sqrt(a), sqrt(1-a))
        │
        ▼
[Stage 3: 5-Factor Weighted Optimization Function] -> Match = 100 * sum(w_k * f_k)
        │
        ▼
[Stage 4: Algorithmic Ranking & Priority Dispatch] -> Outputs Recommended Officers with Score %""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("5-Factor Optimization Formula: ", "Score = 100 * [ 0.35(Specialization) + 0.25(Workload Availability: 1 - (load/max)^1.5) + 0.15(Success Rate) + 0.15(Haversine Proximity: exp(-d/100km)) + 0.10(Experience) ].")
    add_bullet("Haversine Distance Metric: ", "Computes great-circle distance between case crime coordinates and officer precinct base using spherical trigonometry (R = 6371 km).")
    add_bullet("Predictive Duration Velocity: ", "Weibull survival curve modeling estimating days-to-close based on evidence complexity and officer velocity.")
    add_script_box("Judges, our Smart Case Assignment engine runs Agent Nexus-Decision to eliminate investigator burnout and bottlenecking. It computes an objective 5-factor optimization matrix that evaluates domain specialization, non-linear caseload capacity, historical clearance rate, Haversine geographic proximity, and field seniority. This guarantees that critical homicides or cyber cases are dispatched to the exact optimal investigator with accurate resolution day estimations.")

    # Section 5: Evidence Management
    add_section_header("SECTION 5: Evidence Management & Degradation Kinetics", "client/pages/Evidence.tsx", "Agent Nexus-Decision (Degradation Engine)")
    add_code_pipeline("""[Evidence Item Stored: Blood, DNA, Ballistics, Mobile Device, Weapon]
        │
        ▼
[Stage 1: Sample Baseline Half-Life Profiling] -> Blood k0=0.042/day, DNA k0=0.015/day
        │
        ▼
[Stage 2: Thermodynamic Arrhenius Reaction Multiplier] -> k_eff = k0 * exp( (Ea/R)*(1/Tref - 1/Tstor) ) * (1 + γ*ΔH)
        │
        ▼
[Stage 3: First-Order Viability Decay Curve] -> V(t) = V0 * exp(-k_eff * t) => Status Badge (Optimal / Degrading / Critical)
        │
        ▼
[Stage 4: 2D-DCT Perceptual Deduplication Hash] -> 64-bit pHash fingerprint => Hamming Distance <= 5 indicates duplicate""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("Arrhenius Degradation Model: ", "Predicts biological decay rate based on storage temperature T and humidity H relative to activation energy Ea = 85 kJ/mol. V(t) >= 75% (Optimal), 40-75% (Degrading), <40% (Critical).")
    add_bullet("DCT Perceptual Hashing (pHash): ", "Computes 2D Discrete Cosine Transform on 32x32 greyscale evidence frames, extracts 8x8 low-frequency matrix, and compares 64-bit median hash fingerprints across open cases.")
    add_script_box("Judges, in our Evidence module, Agent Nexus-Decision models the physical and biochemical decay of perishable evidence. We utilize first-order Arrhenius reaction kinetics to calculate sample viability decay curves based on ambient storage temperature and humidity parameters. Simultaneously, our DCT-based perceptual hashing cross-references evidence signatures across open cases to instantly catch duplicate weapon or image assets.")

    # Section 6: Audit Trail
    add_section_header("SECTION 6: Blockchain Chain of Custody & Audit Trail", "client/pages/AuditTrail.tsx & Audit.tsx", "Agent Crypt-Ledger (Blockchain Verifier)")
    add_code_pipeline("""[System Event Trigger: Evidence Upload / Biometric Scan / Officer Transfer]
        │
        ▼
[Stage 1: Canonical Event Payload Serialization] -> JSON.stringify(sortKeys(payload))
        │
        ▼
[Stage 2: Cryptographic Leaf Hashing] -> H_event = SHA-256( Canonical_String(S) )
        │
        ▼
[Stage 3: Merkle Tree State-Transition Chaining] -> H_N = SHA-256( H_{N-1} || H_event || Timestamp || OfficerID )
        │
        ▼
[Stage 4: Immutable Audit Trail Timeline] -> Green Verified Badges & Non-Repudiation Logs""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("State-Transition Chaining: ", "Every block N is cryptographically dependent on block N-1: H_N = SHA-256( H_{N-1} || H_event || Timestamp || OfficerID ).")
    add_bullet("Tamper-Evident Proof: ", "Modifying any historical record j causes all subsequent hashes H_{j+1} ... H_N to fail validation immediately, flagging the entire ledger as compromised.")
    add_script_box("Judges, courtroom admissibility hinges on an untampered chain of custody. Agent Crypt-Ledger implements an immutable state-transition audit ledger. Every evidence upload, biometric verification, and investigator transfer is canonically serialized and stamped into a cryptographic SHA-256 Merkle chain. Any retroactive modification to historical records instantly breaks the hash cascade, guaranteeing mathematical non-repudiation.")

    # Section 7: Reports & Collaboration
    add_section_header("SECTION 7: Automated Court Dossiers & Live Collaboration", "client/pages/Reports.tsx & Collaboration.tsx", "Agent Crypt-Ledger (Dossier Synthesizer)")
    add_code_pipeline("""[All Multi-Agent Forensic Outputs Across Modules]
        │
        ▼
[Stage 1: Multi-Agent Neural Output Aggregation] -> Compiles 3D Vectors, Heatmaps, Biometrics, Hashes
        │
        ▼
[Stage 2: Juridical Template Chain-of-Thought Synthesis] -> Generates Formal Legal Sections
        │
        ▼
[Stage 3: Interactive Collaborative Canvas] -> Real-Time Peer Node Mapping & Evidence Tagging
        │
        ▼
[Stage 4: Multi-Signature Approval & PDF Export] -> Cryptographic Sign-Off from Lead Officer & Director""")
    add_subheading("Algorithmic Technicalities & Math:")
    add_bullet("Multi-Agent Dossier Assembly: ", "Standardizes 3D spatial trajectory, deepfake heatmaps, biometric confidence scores, and Merkle block hashes into legal PDF dossiers.")
    add_bullet("Collaborative Multi-Sig Finalization: ", "Interactive investigation canvas requiring dual cryptographic signatures before case files can be submitted to court.")
    add_script_box("Finally, on our Reports and Collaboration tabs, Agent Crypt-Ledger compiles findings from all modules into a court-ready forensic dossier with a single click, while our tactical board allows investigators to collaborate live and seal reports with multi-signature approvals.")

    # Section 8: FAQ Masterclass
    add_subheading("SECTION 8: Judge FAQ & Defense Masterclass")
    faq_table = doc.add_table(rows=5, cols=2)
    faq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    faq_table.autofit = False
    faq_table.columns[0].width = Inches(2.2)
    faq_table.columns[1].width = Inches(4.8)

    faq_hdr = ["Judge Question", "Winning Technical Defense"]
    for i, h in enumerate(faq_hdr):
        cell = faq_table.rows[0].cells[i]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    faq_data = [
        ("How does the 3D Scene generator handle ambiguous prompts?", "Agent Apex-Vision uses deterministic spatial rule priors. If a prompt omits a specific wall, our spatial heuristic defaults to the primary lighting axis (the North wall) and applies a seeded pseudo-random offset within bounded coordinate constraints."),
        ("Why is Grad-CAM better than standard saliency maps in Deepfake detection?", "Standard saliency maps compute simple pixel gradients which are noisy and unspecific. Grad-CAM uses the gradient of the fake-class logit with respect to the final convolutional feature maps, applying a ReLU activation to exclusively highlight features that actively contribute to the synthetic classification."),
        ("How does your Biometric matching handle dirty or smudged latent prints?", "Rather than pixel-level template matching, Agent Bio-Topology abstracts fingerprints into non-Euclidean topological graphs of ridge bifurcations. Because graph adjacency and relative Delaunay triangulation angles are preserved even when 60% of the print is smudged, our Hungarian bipartite matcher maintains high-precision identification."),
        ("What prevents an insider database administrator from falsifying audit records?", "Our Audit Trail uses backward-chained SHA-256 Merkle hashes. If an administrator alters a database row, all downstream block hashes in the chain break immediately, failing verification and alerting the court.")
    ]

    for row_idx, (q, a) in enumerate(faq_data):
        row_cells = faq_table.rows[row_idx + 1].cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate([q, a]):
            set_cell_background(row_cells[col_idx], bg_col)
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8)
            r.font.color.rgb = DARK_TEXT

    # Save output
    output_path = r"c:\Users\manoj\Downloads\orbit-oasis (1)\Orbit_Oasis_Deep_Dive_Tech_Pipelines.docx"
    doc.save(output_path)
    print(f"Deep-Dive Technical Word Document saved to: {output_path}")

if __name__ == "__main__":
    generate_pipeline_docx()
