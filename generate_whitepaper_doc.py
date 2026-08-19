import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(40, 40, 40)

    # Palette
    PRIMARY_COLOR = RGBColor(15, 23, 42)      # Deep Slate / Navy
    SECONDARY_COLOR = RGBColor(14, 116, 144)  # Cyan / Teal
    ACCENT_COLOR = RGBColor(99, 102, 241)     # Indigo / Violet
    DARK_TEXT = RGBColor(30, 41, 59)
    MUTED_TEXT = RGBColor(100, 116, 139)

    # Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("ORBIT-OASIS: TECHNICAL WHITEPAPER")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Decentralized Autonomous Multi-Agent Swarm (MAS), Multi-Modal GraphRAG, and Zero-Knowledge (zk-SNARK) Neural Forensic Verification Protocol")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    # Meta banner table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    col_widths = [Inches(2.3), Inches(2.3), Inches(2.3)]
    for i, w in enumerate(col_widths):
        meta_table.columns[i].width = w

    meta_items = [
        ("Architecture Version", "v4.8 Enterprise-Edge"),
        ("Cryptographic Standard", "zk-SNARK / Groth16 / EVM"),
        ("Consensus Paradigm", "BFT Bayesian Swarm Mesh")
    ]

    for i, (k, v) in enumerate(meta_items):
        cell = meta_table.cell(0, i)
        cell.width = col_widths[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{k}\n")
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = MUTED_TEXT
        r2 = p.add_run(v)
        r2.font.size = Pt(9.5)
        r2.font.bold = True
        r2.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = SECONDARY_COLOR
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = ACCENT_COLOR
        return h

    def add_body(text, space_after=6, italic=False, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK_TEXT
        r.font.italic = italic
        r.font.bold = bold
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.bold = True
        r_prefix.font.color.rgb = DARK_TEXT
        r_text = p.add_run(text)
        r_text.font.color.rgb = DARK_TEXT
        return p

    def add_callout(title, body):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.9)
        cell = table.cell(0, 0)
        set_cell_background(cell, "EEF2FF")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"🔒 {title}\n")
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = ACCENT_COLOR
        
        r2 = p.add_run(body)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Executive Abstract
    add_heading_1("1. Executive Abstract & High-Level Philosophy")
    add_body(
        "Orbit-Oasis introduces a zero-trust, multi-spectral forensic verification paradigm designed to address systemic vulnerabilities in digital asset originality, academic integrity, and biometric credential authentication. Modern generative models (Diffusion Models, Generative Adversarial Networks, and Large Language Models) have created unprecedented capability for high-fidelity forgery, automated plagiarism, and biometric spoofing. "
        "Orbit-Oasis resolves this crisis by unifying three state-of-the-art computational layers into an interdependent trust framework: (1) an Autonomous Multi-Agent Swarm for specialized multi-modal forensics, (2) a Graph-Augmented Retrieval (GraphRAG) architecture for deep semantic lineage cross-matching, and (3) a Zero-Knowledge EVM-anchored Blockchain settlement layer for mathematically tamper-proof audit trails."
    )

    # 2. Autonomous Multi-Agent Swarm (MAS) Engine
    add_heading_1("2. Autonomous Multi-Agent Swarm (MAS) Architecture")
    add_body(
        "Rather than relying on monolithic, single-point-of-failure neural networks, Orbit-Oasis implements an asynchronous, distributed Multi-Agent Swarm (MAS) operating across specialized containerized micro-agents. Communication occurs over high-throughput gRPC message streams with bidirectional actor-critic validation."
    )
    
    add_bullet("Agent-Alpha (Frequency-Domain Deepfake Analyst): ", "Applies 2D Fast Fourier Transforms (2D-FFT) and Discrete Cosine Transforms (DCT) across spatial-temporal video slices. It evaluates azimuthal spectral distribution anomalies to isolate GAN up-sampling checkerboard artifacts and diffusion-based phase irregularities in sub-pixel domains.")
    add_bullet("Agent-Beta (Biometric Topological Graph Neural Network): ", "Deconstructs iris patterns and fingerprint minutiae into dynamic Non-Euclidean geometric graphs. Extracts ridge bifurcation nodes, core delta orientations, and fractional pupil-to-iris pupillary dynamics, comparing them against localized edge-encrypted feature embeddings.")
    add_bullet("Agent-Gamma (GraphRAG Stylometric & Code Analysis Agent): ", "Executes Abstract Syntax Tree (AST) tokenization, semantic n-gram perplexity profiling, and token burstiness entropy calculations to detect LLM synthetic generation and cross-repository structural plagiarism.")
    add_bullet("Agent-Delta (Byzantine Fault Tolerant Swarm Consensus Orchestrator): ", "Aggregates confidence probability tensors from all agents, computing a weighted Bayesian ensemble score and enforcing a Byzantine Fault Tolerant threshold before issuing signed verification certificates.")

    add_callout(
        "Swarm Consensus Mathematical Formulation",
        "The global authenticity score Φ is computed as: Φ = ∑ (w_i · P_i) / (∑ w_i + λ · H(S)), where P_i is the confidence tensor of agent i, w_i is the dynamic reliability weight calibrated via historical validation epochs, and H(S) represents the cross-entropy variance across the swarm to penalize agent hallucinations."
    )

    # 3. Multi-Modal GraphRAG (Retrieval-Augmented Generation)
    add_heading_1("3. Multi-Modal GraphRAG & Semantic Lineage Engine")
    add_body(
        "Standard vector-only RAG systems suffer from severe contextual blind spots, hallucinations, and inability to trace non-linear code/text transformations. Orbit-Oasis develops a hybrid Sparse-Dense Graph-Augmented Retrieval (GraphRAG) system running on high-dimensional Qdrant vector indices synchronized with a Neo4j Property Knowledge Graph."
    )
    
    add_heading_2("Key Technical Capabilities of GraphRAG in Orbit-Oasis:")
    add_bullet("Hierarchical Semantic Chunking: ", "Deconstructs submissions into hierarchical abstract units: Character-level Token Embeddings, Statement Nodes, and Function/Block Subgraphs, preserving deep architectural intent.")
    add_bullet("Dual Sparse-Dense Hybrid Indexing: ", "Combines dense vector embeddings (OpenAI text-embedding-3-large & custom CodeBERT embeddings with 3072 dimensions) with BM25 sparse keyword representations using Reciprocal Rank Fusion (RRF).")
    add_bullet("Citation & Provenance Knowledge Graph: ", "Maintains persistent graph relationships (e.g., [DerivedFrom], [SyntacticallyMutated], [CrossLingualMapped]), allowing the system to track whether a code segment was translated from C++ to Python or synthesized via automated prompt chains.")

    # 4. Zero-Knowledge Cryptography & Blockchain Settlement
    add_heading_1("4. Zero-Knowledge Cryptography & Decentralized Ledger Layer")
    add_body(
        "Privacy and immutable verifiability represent the core pillars of Orbit-Oasis. All forensic conclusions are permanently anchored to an EVM-compatible decentralized ledger via Zero-Knowledge Succinct Non-Interactive Arguments of Knowledge (zk-SNARKs)."
    )

    add_bullet("Zero-Knowledge Verification (zk-SNARKs / Groth16): ", "Enables students, forensic investigators, and institutions to mathematically prove that a digital asset scored above compliance thresholds without revealing the raw asset, biometric data, or proprietary source code to the public ledger.")
    add_bullet("On-Chain Smart Contract Registry: ", "Solidity smart contracts deployed on Ethereum Layer-2 (Arbitrum / Polygon) maintain an immutable, tamper-evident registry of Merkle root hashes, cryptographic timestamps, and zero-knowledge verification proofs.")
    add_bullet("Decentralized Permanent Storage (IPFS / Arweave): ", "Sanitized, cryptographically sealed verification dossiers are uploaded to decentralized content-addressable storage networks, ensuring permanent auditability resistant to centralized server outages or malicious tampering.")

    # Table: Architectural Layer Comparison
    add_heading_2("Table 1: System Architectural Stack & Layer Specifications")
    
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Layer", "Core Technology", "Primary Protocol / Model", "Security / Performance Guarantee"]
    widths = [Inches(1.2), Inches(1.8), Inches(2.2), Inches(1.7)]
    
    for i in range(4):
        table.columns[i].width = widths[i]

    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("Agent Swarm", "Autonomous AI Agents", "gRPC Mesh + Bayesian BFT Ensemble", "Multi-Modal Byzantine Fault Tolerance"),
        ("GraphRAG", "Qdrant + Neo4j Graph", "Sparse-Dense Hybrid HNSW + AST", "Sub-15ms Exact Semantic Origin Tracing"),
        ("Neural Vision", "2D-FFT / DCT Spectral", "ResNet-ST + Minutiae GNN", "Sub-pixel Artifact & Phase Detection"),
        ("Blockchain", "zk-SNARKs + EVM Smart Contracts", "Groth16 Verifier + ERC-4337 + IPFS", "Zero Data Exposure & 100% Immutability")
    ]

    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. End-to-End Operational Lifecycle
    add_heading_1("5. End-to-End Operational Verification Lifecycle")
    add_body("The complete verification lifecycle within Orbit-Oasis follows a strict 6-phase cryptographic pipeline:")
    
    add_bullet("Phase 1 - Ingestion & Edge Pre-Hashing: ", "Input assets (media, code, biometric frames) are received by client WASM workers. Irreversible Poseidon hashes are computed client-side.")
    add_bullet("Phase 2 - Parallelized Swarm Dispatch: ", "The Consensus Orchestrator routes the asset tensors concurrently to Agent-Alpha (Deepfake), Agent-Beta (Biometrics), and Agent-Gamma (GraphRAG).")
    add_bullet("Phase 3 - Multi-Modal Forensic Extraction: ", "Spectral decomposition, Minutiae topological graph matching, and dynamic semantic vector lookups occur within dedicated microsecond compute pipelines.")
    add_bullet("Phase 4 - Swarm Consensus & Calibration: ", "Agent tensors are cross-validated against the confidence threshold; anomaly deviations trigger iterative re-evaluation loops.")
    add_bullet("Phase 5 - Zero-Knowledge Proof Synthesis: ", "The zk-Engine creates a Groth16 zk-SNARK proof attesting to the validity of the verification score without revealing sensitive raw data.")
    add_bullet("Phase 6 - On-Chain Settlement & IPFS Archival: ", "The verification hash is broadcasted to the EVM smart contract and stored permanently on IPFS with an unforgeable cryptographic stamp.")

    # 6. Judge Pitching & Defense Playbook
    add_heading_1("6. Comprehensive Hackathon / Investor Pitch & Defense Playbook")
    add_body(
        "When defending the platform in front of technical judges, investors, or evaluation committees, utilize the following precise framing and structured responses."
    )

    add_heading_2("30-Second Elevator Pitch")
    add_body(
        "\"Orbit-Oasis is a decentralized forensic intelligence platform that combines Autonomous Multi-Agent Swarm intelligence, Graph-Augmented Retrieval (GraphRAG), and Zero-Knowledge Blockchain verification. We provide mathematically tamper-proof, privacy-preserving validation for code originality, biometric identity, and deepfake detection with zero centralized failure points.\"",
        italic=True,
        bold=True
    )

    add_heading_2("Key Technical Defense Strategies (Q&A)")

    add_bullet("Q: 'Why is Blockchain necessary instead of a standard PostgreSQL database?'", "")
    add_body("Defense: Centralized databases have single-point-of-failure vulnerabilities and can be altered by compromised administrators. In high-stakes academic, forensic, and legal environments, proof of non-tampering is paramount. We anchor cryptographic Merkle roots of our forensic audits onto an immutable EVM ledger. Using zk-SNARKs, institutions can verify compliance without ever viewing confidential student submissions or biometric templates.")

    add_bullet("Q: 'How does your Multi-Agent Swarm prevent conflicting detection results?'", "")
    add_body("Defense: We utilize a Byzantine Fault Tolerant (BFT) Bayesian Orchestrator. Each agent outputs a multi-dimensional confidence tensor rather than a binary flag. The Orchestrator computes an entropy-weighted convergence matrix that penalizes outlier uncertainty and reaches verifiable mathematical consensus across all forensic modalities.")

    add_bullet("Q: 'How does GraphRAG outperform traditional Vector Similarity RAG?'", "")
    add_body("Defense: Standard vector RAG computes simple cosine similarity between text chunks, missing structural mutations, paraphrasing chains, and cross-language translation. GraphRAG builds an explicit Knowledge Graph of code syntax trees and semantic entities alongside vector embeddings, enabling multi-hop lineage detection that exposes even heavily obfuscated synthetic plagiarism.")

    add_bullet("Q: 'How do you preserve biometric privacy under strict regulatory frameworks (GDPR / HIPAA)?'", "")
    add_body("Defense: Raw biometric images (iris or fingerprint scans) are never transmitted or stored on centralized servers. Feature extraction occurs in client-side WebAssembly enclaves, transforming biological data into irreversible salted Bio-Hashes verified exclusively through Zero-Knowledge commitments.")

    # Save document
    output_path = r"c:\Users\manoj\Downloads\orbit-oasis (1)\Orbit_Oasis_Advanced_Technical_Whitepaper.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_document()
