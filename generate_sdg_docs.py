import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import odf
from odf.opendocument import OpenDocumentText
from odf.text import P, H, Span
from odf.table import Table, TableColumn, TableRow, TableCell
from odf.style import Style, TextProperties, ParagraphProperties, TableProperties, TableColumnProperties, TableCellProperties

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

sdg_data = [
    {
        "goal": "1. No Poverty",
        "ideas": [
            "AI-driven job matching platforms & legal aid accessibility",
            "Blockchain for transparent aid distribution & anti-fraud verification",
            "Mobile apps for financial literacy & micro-finance security",
            "Orbit-Oasis Application: Pro-bono legal & forensic audit access to eliminate socio-economic bias in criminal defense"
        ],
        "tech": [
            "Python, TensorFlow, NLP (for job-matching AI & legal NLP)",
            "Ethereum, Hyperledger Fabric, SHA-256 (for blockchain aid auditing)",
            "Flutter, React.js, Firebase (for accessible mobile apps)"
        ]
    },
    {
        "goal": "2. Zero Hunger",
        "ideas": [
            "Smart farming with IoT sensors & yield prediction",
            "AI for crop disease detection & pest prevention",
            "Supply chain optimization for food distribution & anti-theft tracing",
            "Orbit-Oasis Application: Tamper-proof IoT ledger for agricultural supply chain logistics & spoilage prediction"
        ],
        "tech": [
            "Arduino, Raspberry Pi, IoT sensors (for smart farming)",
            "OpenCV, CNN models, ResNet (for crop analysis)",
            "AWS, Google Cloud AI, Node.js (for supply chain tracking)"
        ]
    },
    {
        "goal": "3. Good Health & Well-being",
        "ideas": [
            "Telemedicine & AI clinical diagnosis platforms",
            "Mental health chatbots & stress monitoring systems",
            "Wearable health telemetry & vitals tracking",
            "Orbit-Oasis Application: Investigator burnout prevention via 5-Factor automated workload balancing & cognitive fatigue minimization"
        ],
        "tech": [
            "TensorFlow, PyTorch (for medical diagnosis)",
            "Wearable APIs (Fitbit, Apple Health)",
            "Dialogflow, GPT APIs, FastAPI (for health chatbots)",
            "Optimization Engine: SciPy Hungarian bipartite matching"
        ]
    },
    {
        "goal": "4. Quality Education",
        "ideas": [
            "AI-powered personalized learning & adaptive tutoring",
            "VR/AR for immersive educational simulations",
            "Apps for literacy & numeracy improvement",
            "Orbit-Oasis Application: 3D interactive WebGL crime scene simulator for forensic science education & training cadets"
        ],
        "tech": [
            "Python, Scikit-learn (for adaptive learning)",
            "Three.js, WebGL, Unity, React Three Fiber (for 3D spatial simulations)",
            "Swift, Kotlin, Flutter, React SPA (for educational apps)"
        ]
    },
    {
        "goal": "5. Gender Equality",
        "ideas": [
            "AI bias detection in hiring & evaluation algorithms",
            "Safe anonymous reporting apps for harassment & abuse",
            "Digital mentorship platforms & equal opportunity tracking",
            "Orbit-Oasis Application: Algorithmic fairness in law enforcement dispatch & non-discriminatory suspect biometric matching"
        ],
        "tech": [
            "Google AutoML, Fairness Indicators (for AI bias auditing)",
            "Flutter, Firebase, WebSockets (for encrypted safe reporting)",
            "WebRTC, React.js, Zod (for secure mentorship platforms)"
        ]
    },
    {
        "goal": "6. Clean Water & Sanitation",
        "ideas": [
            "IoT-based water quality & contamination monitoring",
            "AI for pipeline leak detection & sewer management",
            "Blockchain for municipal water resource management",
            "Orbit-Oasis Application: Biochemical degradation modeling (Arrhenius half-life) applicable to environmental contaminant decay"
        ],
        "tech": [
            "Arduino, NodeMCU, IoT chemical sensors (for water quality)",
            "Computer Vision, OpenCV, Edge AI (for leak detection)",
            "Ethereum, Smart Contracts, PostgreSQL (for water governance)"
        ]
    },
    {
        "goal": "7. Affordable & Clean Energy",
        "ideas": [
            "Smart energy grids with AI load balancing",
            "IoT-based industrial energy optimization",
            "Blockchain for decentralized peer-to-peer energy trading",
            "Orbit-Oasis Application: Low-power edge AI inference optimization (ONNX / TensorRT) for green computing in forensic labs"
        ],
        "tech": [
            "MATLAB, TensorFlow, ONNX Runtime (for energy grids & lightweight AI)",
            "NodeMCU, ESP32, MQTT (for energy monitoring)",
            "Solidity, Ethereum, Web3.js (for decentralized energy trading)"
        ]
    },
    {
        "goal": "8. Decent Work & Economic Growth",
        "ideas": [
            "AI-driven job market analytics & career pathing",
            "Virtual skill development & workforce training platforms",
            "Crowdsourcing & ethical gig economy job portals",
            "Orbit-Oasis Application: Intelligent police resource dispatching to optimize officer shift capacity & eliminate administrative overtime"
        ],
        "tech": [
            "BigQuery, Tableau, PostgreSQL (for labor analytics)",
            "WebRTC, React.js, TailwindCSS (for training platforms)",
            "FastAPI, Flask, Express.js (for workflow dispatch engines)"
        ]
    },
    {
        "goal": "9. Industry, Innovation & Infrastructure",
        "ideas": [
            "Smart traffic management & urban mobility AI",
            "AI-driven predictive maintenance for critical infrastructure",
            "Digital twins & 3D procedural simulation for structural planning",
            "Orbit-Oasis Application: Full-scale 3D procedural crime scene spatial reconstruction with inverse raycasting laser trajectories"
        ],
        "tech": [
            "Three.js, WebGL 2.0, React Three Fiber, Drei (for 3D digital twins)",
            "YOLO, OpenCV, PyTorch (for spatial video inspection)",
            "IoT, Edge Computing, Docker (for resilient infrastructure)"
        ]
    },
    {
        "goal": "10. Reduced Inequalities",
        "ideas": [
            "AI-driven accessibility tools for disabled individuals",
            "Real-time multilingual translation apps for minority groups",
            "Blockchain for fair wages & uncorrupted registry tracking",
            "Orbit-Oasis Application: Explainable AI (XAI with Grad-CAM heatmaps & Delaunay biometric graphs) to prevent wrongful arrests and discrimination"
        ],
        "tech": [
            "AI4Accessibility, TensorFlow, Grad-CAM (for explainable AI & accessibility)",
            "Google Translate API, Whisper, NLP (for real-time translation)",
            "Hyperledger, Smart Contracts (for transparent public registries)"
        ]
    },
    {
        "goal": "11. Sustainable Cities & Communities",
        "ideas": [
            "Smart municipal waste management & bin sensors",
            "AI for urban air pollution monitoring & hazard alerting",
            "Urban mobility & emergency response route optimization",
            "Orbit-Oasis Application: Geospatial forensic mapping & crime hotspot cluster analysis for safer urban neighborhoods"
        ],
        "tech": [
            "IoT Smart Sensors, ESP32 (for waste & pollution monitoring)",
            "Computer Vision, Deep Learning, GIS (for hazard tracking)",
            "Graph Algorithms (Dijkstra, A*), Leaflet.js, Turf.js (for emergency dispatch)"
        ]
    },
    {
        "goal": "12. Responsible Consumption & Production",
        "ideas": [
            "AI for automated industrial waste sorting",
            "Blockchain for ethical supply chain sourcing & traceability",
            "Smart inventory & perishable item tracking",
            "Orbit-Oasis Application: Arrhenius bio-sample decay tracking to eliminate forensic evidence spoilage and chemical reagent waste"
        ],
        "tech": [
            "Computer Vision, OpenCV, Roboflow (for waste sorting)",
            "Ethereum, Hyperledger, IPFS (for ethical supply chains)",
            "SciPy Arrhenius equations, AWS, Google Cloud AI (for inventory tracking)"
        ]
    },
    {
        "goal": "13. Climate Action",
        "ideas": [
            "AI for climate prediction & natural disaster alerting",
            "Carbon footprint calculation & corporate offset tracking apps",
            "IoT-based forest fire and environmental monitoring mesh",
            "Orbit-Oasis Application: Paperless, end-to-end digital juridical reporting replacing physical case binders with digital SHA-256 dossiers"
        ],
        "tech": [
            "NASA Earth Data, Machine Learning, Scikit-Learn (for climate modeling)",
            "Flutter, React.js, Firebase (for carbon tracking)",
            "IoT, Edge Computing, LoRaWAN (for disaster alert meshes)"
        ]
    },
    {
        "goal": "14. Life Below Water",
        "ideas": [
            "AI for marine species classification & coral reef health monitoring",
            "IoT acoustic & water sensors for illegal ocean dumping",
            "Blockchain for sustainable fishing & quota compliance",
            "Orbit-Oasis Application: Dual neural network feature extraction adaptable to underwater forensic search and sonar imagery"
        ],
        "tech": [
            "TensorFlow, Computer Vision, ResNet (for marine bio-monitoring)",
            "IoT Water Sensors, Underwater Acoustic Modems (for ocean sensors)",
            "Solidity, Hyperledger Fabric (for fishery compliance logs)"
        ]
    },
    {
        "goal": "15. Life on Land",
        "ideas": [
            "AI for illegal deforestation detection via satellite imagery",
            "Wildlife tracking & anti-poaching surveillance with drones",
            "Smart precision irrigation with IoT soil probes",
            "Orbit-Oasis Application: Forensic optical verification for wildlife crime tracking and anti-poaching biometric tagging"
        ],
        "tech": [
            "Satellite Data (Sentinel/Landsat), Deep Learning (for forest monitoring)",
            "Drones, OpenCV, YOLOv8 (for anti-poaching thermal tracking)",
            "Raspberry Pi, IoT Soil Sensors (for smart irrigation)"
        ]
    },
    {
        "goal": "16. Peace, Justice & Strong Institutions (CORE DIRECT ALIGNMENT)",
        "ideas": [
            "Orbit-Oasis Core Platform: Multi-Agent AI Forensic Intelligence & Evidence Integrity Mesh",
            "Digital Forensics & Deepfake Detection: Dual-backbone neural ensemble to expose synthetic media, fraudulent testimonies, and manipulated footage",
            "Cryptographic Chain of Custody: SHA-256 Merkle block ledger guaranteeing 100% tamper-evident evidence auditing from crime scene to court",
            "Multi-Modal Biometrics: 2048-bit Daugman iris wavelets + Hungarian graph minutiae matching for rapid, unbiased suspect resolution in <50ms",
            "Automated Juridical Report Synthesis: Instant court-admissible legal dossiers with transparent XAI visual artifacts (Grad-CAM heatmaps)",
            "Predictive Policing & Case Velocity: Reducing multi-month forensic backlogs and preventing human error in judicial systems"
        ],
        "tech": [
            "PyTorch, Torchvision, Dual CNN-ViT Ensembles (for deepfake & media forensics)",
            "OpenCV, Daugman Wavelet Filters, SciPy Hungarian Algorithm (for multi-modal biometrics)",
            "Node.js Crypto, SHA-256 Merkle Trees, IPFS / Pinata (for tamper-proof chain of custody)",
            "Three.js, React Three Fiber, Drei, WebGL (for 3D spatial crime scene laser trajectory solver)",
            "Express.js, TypeScript, React 18 SPA, TailwindCSS, PostgreSQL (for full-stack forensic hub)"
        ]
    },
    {
        "goal": "17. Partnerships for the Goals",
        "ideas": [
            "Open data platforms for cross-jurisdictional SDG tracking & metrics",
            "AI for UN sustainability reporting & automated compliance audits",
            "Cross-border inter-agency forensic data sharing & consortium networks",
            "Orbit-Oasis Application: Standardized CJIS API endpoints & IPFS decentralized storage enabling global law enforcement collaboration"
        ],
        "tech": [
            "Google Data Studio, BigQuery, REST APIs (for global SDG tracking)",
            "NLP, LLM fine-tuning, Juridical Report Generators (for compliance reporting)",
            "React.js, Node.js, Express, IPFS, Docker (for federated inter-agency networks)"
        ]
    }
]

def generate_word_doc(output_path):
    doc = docx.Document()
    
    # Page setup - 0.6 inch margins for table space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("UN Sustainable Development Goals (SDGs) & Computer Science Project Mapping")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42) # Dark Slate
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(14)
    subtitle_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Applied Technical Matrix Featuring the Orbit-Oasis Forensic AI & Evidence Integrity Platform")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600

    # Table creation
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Column widths (Total = 7.3 inches)
    col_widths = [Inches(1.8), Inches(3.2), Inches(2.3)]
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["SDG Goal", "CS Project Ideas & Project Integration", "Tech Tools & Technologies"]
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, "0F172A") # Deep Navy
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for idx, item in enumerate(sdg_data):
        row = table.add_row()
        for i, w in enumerate(col_widths):
            row.cells[i].width = w
            
        c_goal, c_ideas, c_tech = row.cells
        
        # Highlight SDG 16 or alternate shading
        is_highlight = "16" in item["goal"]
        bg_color = "F1F5F9" if is_highlight else ("FAFAFA" if idx % 2 == 1 else "FFFFFF")
        
        for c in [c_goal, c_ideas, c_tech]:
            set_cell_background(c, bg_color)
            set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        
        # 1. Goal Cell
        p_g = c_goal.paragraphs[0]
        p_g.paragraph_format.space_before = Pt(0)
        p_g.paragraph_format.space_after = Pt(0)
        r_g = p_g.add_run(item["goal"])
        r_g.font.name = "Arial"
        r_g.font.size = Pt(9.5)
        r_g.font.bold = True
        if is_highlight:
            r_g.font.color.rgb = RGBColor(3, 105, 161) # Highlight blue
        else:
            r_g.font.color.rgb = RGBColor(15, 23, 42)
            
        # 2. Ideas Cell
        c_ideas.paragraphs[0].text = "" # Clear default
        for j, idea in enumerate(item["ideas"]):
            p_i = c_ideas.add_paragraph() if j > 0 else c_ideas.paragraphs[0]
            p_i.paragraph_format.space_before = Pt(0)
            p_i.paragraph_format.space_after = Pt(2)
            p_i.paragraph_format.left_indent = Inches(0.12)
            
            # Format bullet
            if idea.startswith("Orbit-Oasis"):
                r_bullet = p_i.add_run("★ ")
                r_bullet.font.bold = True
                r_bullet.font.color.rgb = RGBColor(14, 116, 144)
                r_text = p_i.add_run(idea)
                r_text.font.bold = True
                r_text.font.size = Pt(8.5)
                r_text.font.color.rgb = RGBColor(15, 23, 42)
            else:
                r_bullet = p_i.add_run("- ")
                r_bullet.font.bold = True
                r_text = p_i.add_run(idea)
                r_text.font.size = Pt(8.5)
                r_text.font.color.rgb = RGBColor(51, 65, 85)
            r_text.font.name = "Arial"
            
        # 3. Tech Cell
        c_tech.paragraphs[0].text = ""
        for k, tech in enumerate(item["tech"]):
            p_t = c_tech.add_paragraph() if k > 0 else c_tech.paragraphs[0]
            p_t.paragraph_format.space_before = Pt(0)
            p_t.paragraph_format.space_after = Pt(2)
            p_t.paragraph_format.left_indent = Inches(0.12)
            
            r_bullet = p_t.add_run("- ")
            r_bullet.font.bold = True
            r_text = p_t.add_run(tech)
            r_text.font.name = "Arial"
            r_text.font.size = Pt(8.5)
            r_text.font.color.rgb = RGBColor(30, 41, 59)

    set_table_borders(table, color="CBD5E1", sz="4", val="single")

    # Add summary section at the bottom
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Executive Alignment Summary: Orbit-Oasis Platform")
    h2_run.font.name = "Arial"
    h2_run.font.size = Pt(12)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(15, 23, 42)
    
    summary_text = (
        "The Orbit-Oasis project primarily anchors on UN SDG 16 (Peace, Justice and Strong Institutions) "
        "by introducing an automated multi-agent AI forensic ecosystem that tackles case backlogs, eliminates evidence "
        "tampering through SHA-256 cryptographic chain-of-custody ledgers, provides explainable deepfake and biometric "
        "authenticity verification, and ensures transparent, fair judicial outcomes. Secondary alignment spans SDG 9 "
        "(Industry, Innovation & Infrastructure via 3D procedural raycast modeling), SDG 8 (Decent Work through AI caseload balancing), "
        "SDG 10 (Reduced Inequalities via unbiased XAI algorithms), and SDG 17 (Global Inter-Agency Forensic Standards)."
    )
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(6)
    r_sum = p_sum.add_run(summary_text)
    r_sum.font.name = "Arial"
    r_sum.font.size = Pt(9.5)
    r_sum.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    print(f"Successfully generated Word Document: {output_path}")

def generate_odt_doc(output_path):
    odt_doc = OpenDocumentText()
    
    # Define styles
    title_style = Style(name="DocTitle", family="paragraph")
    title_style.addElement(ParagraphProperties(textalign="center", margintop="0pt", marginbottom="6pt"))
    title_style.addElement(TextProperties(fontsize="16pt", fontweight="bold", fontfamily="Arial", color="#0f172a"))
    odt_doc.styles.addElement(title_style)
    
    subtitle_style = Style(name="DocSubtitle", family="paragraph")
    subtitle_style.addElement(ParagraphProperties(textalign="center", margintop="0pt", marginbottom="14pt"))
    subtitle_style.addElement(TextProperties(fontsize="11pt", fontstyle="italic", fontfamily="Arial", color="#475569"))
    odt_doc.styles.addElement(subtitle_style)
    
    th_style = Style(name="TableHeader", family="table-cell")
    th_style.addElement(TableCellProperties(backgroundcolor="#0f172a", padding="6pt", border="0.5pt solid #cbd5e1"))
    odt_doc.styles.addElement(th_style)
    
    th_p_style = Style(name="TableHeaderP", family="paragraph")
    th_p_style.addElement(TextProperties(fontsize="10pt", fontweight="bold", fontfamily="Arial", color="#ffffff"))
    odt_doc.styles.addElement(th_p_style)
    
    cell_style = Style(name="TableCell", family="table-cell")
    cell_style.addElement(TableCellProperties(padding="5pt", border="0.5pt solid #cbd5e1"))
    odt_doc.styles.addElement(cell_style)
    
    cell_hl_style = Style(name="TableCellHL", family="table-cell")
    cell_hl_style.addElement(TableCellProperties(backgroundcolor="#f1f5f9", padding="5pt", border="0.5pt solid #38bdf8"))
    odt_doc.styles.addElement(cell_hl_style)
    
    p_body_style = Style(name="TableBodyP", family="paragraph")
    p_body_style.addElement(ParagraphProperties(marginbottom="2pt", margintop="0pt"))
    p_body_style.addElement(TextProperties(fontsize="8.5pt", fontfamily="Arial", color="#334155"))
    odt_doc.styles.addElement(p_body_style)
    
    p_bold_style = Style(name="TableBodyPBold", family="paragraph")
    p_bold_style.addElement(ParagraphProperties(marginbottom="2pt", margintop="0pt"))
    p_bold_style.addElement(TextProperties(fontsize="9pt", fontweight="bold", fontfamily="Arial", color="#0f172a"))
    odt_doc.styles.addElement(p_bold_style)

    # Content
    h_title = P(stylename=title_style, text="UN Sustainable Development Goals (SDGs) & Computer Science Project Mapping")
    odt_doc.text.addElement(h_title)
    
    h_sub = P(stylename=subtitle_style, text="Applied Technical Matrix Featuring the Orbit-Oasis Forensic AI & Evidence Integrity Platform")
    odt_doc.text.addElement(h_sub)
    
    table = Table()
    table.addElement(TableColumn(numbercolumnsrepeated=1))
    table.addElement(TableColumn(numbercolumnsrepeated=1))
    table.addElement(TableColumn(numbercolumnsrepeated=1))
    
    # Header
    hdr_row = TableRow()
    for t in ["SDG Goal", "CS Project Ideas & Project Integration", "Tech Tools & Technologies"]:
        c = TableCell(stylename=th_style)
        c.addElement(P(stylename=th_p_style, text=t))
        hdr_row.addElement(c)
    table.addElement(hdr_row)
    
    # Rows
    for item in sdg_data:
        row = TableRow()
        is_hl = "16" in item["goal"]
        c_style = cell_hl_style if is_hl else cell_style
        
        # Goal
        cg = TableCell(stylename=c_style)
        cg.addElement(P(stylename=p_bold_style, text=item["goal"]))
        row.addElement(cg)
        
        # Ideas
        ci = TableCell(stylename=c_style)
        for idea in item["ideas"]:
            ci.addElement(P(stylename=p_body_style, text=f"- {idea}"))
        row.addElement(ci)
        
        # Tech
        ct = TableCell(stylename=c_style)
        for tech in item["tech"]:
            ct.addElement(P(stylename=p_body_style, text=f"- {tech}"))
        row.addElement(ct)
        
        table.addElement(row)
        
    odt_doc.text.addElement(table)
    odt_doc.save(output_path)
    print(f"Successfully generated ODT Document: {output_path}")

def generate_markdown_doc(output_path):
    lines = [
        "# 🌐 UN Sustainable Development Goals (SDGs) & Computer Science Project Mapping",
        "## *Technical Matrix & Alignment for Orbit-Oasis: Multi-Agent Forensic AI Platform*",
        "",
        "---",
        "",
        "| SDG Goal | CS Project Ideas & Project Integration | Tech Tools & Technologies |",
        "| :--- | :--- | :--- |"
    ]
    
    for item in sdg_data:
        ideas_str = "<br>".join([f"• **{i}**" if i.startswith("Orbit-Oasis") else f"• {i}" for i in item["ideas"]])
        tech_str = "<br>".join([f"• {t}" for t in item["tech"]])
        goal_str = f"**{item['goal']}**" if "16" not in item["goal"] else f"🎯 **{item['goal']}** *(Primary Focus)*"
        lines.append(f"| {goal_str} | {ideas_str} | {tech_str} |")
        
    lines.extend([
        "",
        "---",
        "",
        "## 🏛️ Direct Alignment: Orbit-Oasis Project Breakdown",
        "",
        "| Orbit-Oasis Module / Agent | Core Capability | Targeted UN SDG | Specific UN Indicator & Metric |",
        "| :--- | :--- | :--- | :--- |",
        "| **`Agent Apex-Vision`** (Media Forensics) | Dual-backbone neural network deepfake detection & Grad-CAM visual heatmaps | **SDG 16 (Peace & Justice)**<br>**SDG 9 (Innovation)** | **Target 16.3**: Equal access to justice by eliminating fabricated/synthetic evidence; **Target 9.5**: Advanced scientific research and technological capabilities in judicial systems |",
        "| **`Agent Apex-Vision`** (3D Spatial Ballistics) | Procedural 3D WebGL crime scene reconstruction with inverse raycast trajectory solver | **SDG 16 (Peace & Justice)**<br>**SDG 11 (Safe Cities)** | **Target 16.1**: Reduce violence rates through rigorous ballistic crime scene analysis; **Target 11.7**: Enhancing urban safety & community security |",
        "| **`Agent Bio-Topology`** (Biometric Engine) | 2048-bit Daugman iris wavelet demodulation + Delaunay graph minutiae Hungarian matching | **SDG 16 (Peace & Justice)**<br>**SDG 10 (Reduced Inequalities)** | **Target 16.9**: Legal identity verification; **Target 10.3**: Eliminating wrongful accusations and racial/demographic bias through mathematically proven XAI minutiae vector graphs |",
        "| **`Agent Nexus-Decision`** (Resource & Decay) | 5-Factor officer workload matching & Arrhenius bio-sample half-life degradation simulation | **SDG 8 (Decent Work)**<br>**SDG 12 (Responsible Resource Use)** | **Target 8.8**: Safe, balanced working environments preventing investigator cognitive fatigue; **Target 12.2**: Eliminating spoilage and waste of biochemical evidence |",
        "| **`Agent Crypt-Ledger`** (Audit Trail) | SHA-256 state-transition Merkle tree blockchain with court-admissible NLP legal dossiers | **SDG 16 (Peace & Justice)**<br>**SDG 17 (Partnerships)** | **Target 16.5 & 16.6**: Substantially reduce corruption, evidence tampering, and foster accountable institutions; **Target 17.16**: Global inter-agency CJIS data sharing |",
        "",
        "---",
        ""
    ])
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Successfully generated Markdown Document: {output_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(base_dir, "Orbit_Oasis_SDG_Goals_Mapping.docx")
    odt_path = os.path.join(base_dir, "Orbit_Oasis_SDG_Goals_Mapping.odt")
    md_path = os.path.join(base_dir, "SDG_GOALS_PROJECT_MAPPING.md")
    
    generate_word_doc(docx_path)
    generate_odt_doc(odt_path)
    generate_markdown_doc(md_path)
