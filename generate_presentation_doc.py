import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_presentation_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    PRIMARY = RGBColor(15, 23, 42)       # Slate (#0F172A)
    ACCENT_BLUE = RGBColor(14, 116, 144) # Cyan/Teal (#0E7490)
    ACCENT_PURPLE = RGBColor(99, 102, 241)# Indigo (#6366F1)
    DARK_TEXT = RGBColor(30, 41, 59)
    MUTED_TEXT = RGBColor(100, 116, 139)
    GOLD_ACCENT = RGBColor(180, 83, 9)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(2)
    r_main = title_p.add_run("ORBIT-OASIS: NEXT-GEN FORENSIC AI PLATFORM")
    r_main.font.size = Pt(22)
    r_main.font.bold = True
    r_main.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    r_sub = sub_p.add_run("Executive Pitch Deck & Feature-by-Feature Demonstration Guide")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = ACCENT_BLUE

    def add_slide(slide_num, title, screen_route, agent_name, bullets, speaker_notes, demo_action):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        cell = table.cell(0, 0)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=130, bottom=130, left=160, right=160)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_num = p.add_run(f"SLIDE {slide_num:02d} | ")
        r_num.font.size = Pt(9.5)
        r_num.font.bold = True
        r_num.font.color.rgb = ACCENT_PURPLE
        
        r_title = p.add_run(title)
        r_title.font.size = Pt(13.5)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(255, 255, 255)

        p_sub = cell.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(2)
        p_sub.paragraph_format.space_after = Pt(0)
        r_s = p_sub.add_run(f"🖥️ Screen: {screen_route}  |  🤖 Powered by: {agent_name}")
        r_s.font.size = Pt(9)
        r_s.font.bold = True
        r_s.font.color.rgb = RGBColor(147, 197, 253)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        for bold_prefix, text in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(3)
            bp.paragraph_format.line_spacing = 1.15
            rb = bp.add_run(bold_prefix)
            rb.font.size = Pt(9.5)
            rb.font.bold = True
            rb.font.color.rgb = DARK_TEXT
            rt = bp.add_run(text)
            rt.font.size = Pt(9.5)
            rt.font.color.rgb = DARK_TEXT

        # Live Demo Action Box
        v_table = doc.add_table(rows=1, cols=1)
        v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        v_table.autofit = False
        v_table.columns[0].width = Inches(7.0)
        v_cell = v_table.cell(0, 0)
        set_cell_background(v_cell, "F1F5F9")
        set_cell_margins(v_cell, top=70, bottom=70, left=120, right=120)
        vp = v_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        r_v1 = vp.add_run("👉 What to Click / Show on Website: ")
        r_v1.font.size = Pt(8.5)
        r_v1.font.bold = True
        r_v1.font.color.rgb = ACCENT_BLUE
        r_v2 = vp.add_run(demo_action)
        r_v2.font.size = Pt(8.5)
        r_v2.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

        # Presenter Script
        sn_table = doc.add_table(rows=1, cols=1)
        sn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sn_table.autofit = False
        sn_table.columns[0].width = Inches(7.0)
        sn_cell = sn_table.cell(0, 0)
        set_cell_background(sn_cell, "FEF3C7")
        set_cell_margins(sn_cell, top=70, bottom=70, left=120, right=120)
        snp = sn_cell.paragraphs[0]
        snp.paragraph_format.space_after = Pt(0)
        r_sn_lbl = snp.add_run("🎙️ Presenter Script (Say This): ")
        r_sn_lbl.font.size = Pt(8.5)
        r_sn_lbl.font.bold = True
        r_sn_lbl.font.color.rgb = GOLD_ACCENT
        r_sn = snp.add_run(f'"{speaker_notes}"')
        r_sn.font.size = Pt(9)
        r_sn.font.italic = True
        r_sn.font.color.rgb = DARK_TEXT

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Slide 1: Platform Command Center
    add_slide(
        1,
        "Orbit-Oasis: Intelligent Forensic Department Overview",
        "/dashboard (Main Dashboard)",
        "Unified Multi-Agent Forensic Mesh",
        [
            ("Integrated Forensic Hub: ", "Unifies real-time case analytics, multi-modal evidence intelligence, and blockchain verification in a single command HUD."),
            ("Real-Time Telemetry: ", "Monitors live case resolution velocity, biometric matching queues, and active forensic officer workloads."),
            ("Four Flagship AI Agents: ", "Under the hood, 4 coordinated neural agents process media forensics, biometrics, predictive dispatch, and cryptographic audit trails.")
        ],
        "Good morning judges. Welcome to Orbit-Oasis—an enterprise AI forensic intelligence platform. Right here on our dashboard, our system connects real-time crime scene spatial modeling, deepfake detection, biometric matching, and blockchain chain-of-custody into a unified, high-speed investigation hub.",
        "Show the main Dashboard with active case statistics, evidence metrics, and navigation sidebar."
    )

    # Slide 2: 3D Crime Scene Reconstruction
    add_slide(
        2,
        "3D Volumetric Crime Scene & Spatial Trajectory Modeling",
        "/crime-scene (3D Crime Scene Tab)",
        "Agent Apex-Vision (Spatial Engine)",
        [
            ("Prompt-to-3D Scene Synthesis: ", "Transforms descriptive incident text into interactive 3D WebGL volumetric crime scene environments in real-time."),
            ("Spatial Evidence Mapping: ", "Automatically renders bullet shells, blood spatter, weapons, and points of entry on accurate 3D coordinates."),
            ("Interactive Laser Measurement: ", "Built-in measurement ray-tracing calculates exact distances between evidence markers and suspected shooter origins.")
        ],
        "On our 3D Crime Scene tab, Agent Apex-Vision uses natural language processing to turn written incident descriptions into an interactive 3D volumetric room. Detectives can rotate the scene, inspect blood spatter and bullet casings, and use our spatial measurement tool to trace trajectory lines with millimeter accuracy.",
        "Click the sample prompt buttons (e.g. 'Small office room...'), click 'Generate Scene', rotate the 3D room, and switch tools between 'View', 'Measure', and 'Mark'."
    )

    # Slide 3: Deepfake Detection & Frequency Analysis
    add_slide(
        3,
        "Deepfake Detection & Sub-Pixel Neural Heatmaps",
        "/deepfake (Deepfake Detection Tab)",
        "Agent Apex-Vision (Media Forensics)",
        [
            ("Dual-Model Ensemble Scoring: ", "Combines Primary and Secondary neural vision networks to calculate an ensemble AI-generation confidence score."),
            ("Frame-by-Frame Temporal Scanning: ", "Samples video frames over time, calculating exact fake-frame and AI-frame ratios across the entire clip timeline."),
            ("Grad-CAM Frequency Heatmap: ", "Generates visual activation heatmaps exposing synthetic pixel blending artifacts and unnatural facial seams.")
        ],
        "Here on our Deepfake Detection tab, Agent Apex-Vision tackles synthetic media forgery. It runs a dual-model neural ensemble that evaluates frame-by-frame AI ratios and renders sub-pixel Grad-CAM activation heatmaps, giving investigators visual proof of facial tampering and deepfake synthesis.",
        "Upload a sample image/video, click 'Analyze', and point out the Ensemble Confidence Score, Inference Time (ms), and Grad-CAM Heatmap."
    )

    # Slide 4: Multi-Modal Biometrics (Iris & Fingerprint)
    add_slide(
        4,
        "Multi-Modal Biometrics: Iris & Fingerprint Graph Fusion",
        "/biometric (Biometric Analysis Tab)",
        "Agent Bio-Topology (Biometric Transformer)",
        [
            ("Dual Modality Analysis: ", "Dedicated neural pipelines for high-resolution Iris textural wavelets and Fingerprint ridge minutiae graphs."),
            ("Radial Match Confidence Rings: ", "Visual 0-100% confidence ring scoring with instant classification from 'Partial Match' to 'Match Confirmed'."),
            ("Automated Suspect Cross-Matching: ", "Ranks candidate suspect profiles against criminal history, Aadhaar IDs, age, and location metadata.")
        ],
        "On the Biometric tab, Agent Bio-Topology performs multi-modal matching. It converts fingerprint minutiae into graph adjacency matrices and extracts high-frequency iris patterns. The radial score ring gives instant match verdicts and pulls verified suspect records with full criminal history in under 50 milliseconds.",
        "Toggle between 'Fingerprint' and 'Iris' tabs, upload a scan, and show the glowing radial Score Ring, the match score (e.g. 94%), and the matched Suspect Profile card."
    )

    # Slide 5: Smart Case Assignment & Velocity Prediction
    add_slide(
        5,
        "Predictive Case Velocity & 5-Factor Officer Assignment",
        "/assignment & /cases (Assignment & Cases Tabs)",
        "Agent Nexus-Decision (Predictive Allocator)",
        [
            ("5-Factor Recommendation Engine: ", "Calculates optimal officer matches using Specialization, Workload Availability, Success Rate, Proximity (km), and Experience."),
            ("Predictive Resolution Velocity: ", "Estimates case completion timeline in days based on crime type, evidence complexity, and officer velocity."),
            ("Automated Caseload Balancing: ", "Prevents investigator burnout by dynamically monitoring active caseloads vs. maximum capacity.")
        ],
        "On our Smart Assignment tab, Agent Nexus-Decision solves forensic backlogs. It computes a 5-factor algorithmic match—weighing officer specialization, active caseload, distance, and historical success rate—to assign the perfect investigator and estimate the exact days to resolution.",
        "Navigate to '/assignment', click on a case, and show the 'Recommended Officers' ranking with the match score percentage and breakdown factors."
    )

    # Slide 6: Evidence Degradation & Half-Life Modeling
    add_slide(
        6,
        "Evidence Degradation & Perishable Sample Modeling",
        "/evidence (Evidence Management Tab)",
        "Agent Nexus-Decision (Degradation Engine)",
        [
            ("Perishable Bio-Sample Half-Life: ", "Calculates degradation risk curves for biological materials (blood, DNA, tissue) based on environmental storage conditions."),
            ("Duplicate Evidence Detection: ", "Cross-checks evidence hashes to flag duplicate items and shared forensic signatures across cases."),
            ("Visual Degradation Badges: ", "Color-coded status indicators track evidence viability from 'Optimal' to 'Critical Action Required'.")
        ],
        "Under our Evidence tab, Agent Nexus-Decision models evidence degradation. It computes half-life decay curves for biological samples based on ambient storage parameters, warning officers before crucial DNA or blood samples spoil, while simultaneously flagging duplicate evidence across active cases.",
        "Show the Evidence list, point out the evidence types (Biological, Digital, Physical), and highlight the status tags and degradation parameters."
    )

    # Slide 7: Blockchain Chain of Custody & Audit Trail
    add_slide(
        7,
        "Tamper-Proof Blockchain Audit Trail & zk-Verification",
        "/audit-trail (Audit Trail Tab)",
        "Agent Crypt-Ledger (Blockchain Verifier)",
        [
            ("Immutable Cryptographic Ledger: ", "Logs every single case creation, evidence upload, biometric verification, and officer handoff with timestamped block hashes."),
            ("Verified Action Badges: ", "Color-coded cryptographic audit stamps (e.g. 'Biometric Verified', 'Smart Contract Executed', 'Chain of Custody Transfer')."),
            ("100% Tamper-Evident Security: ", "Guarantees court-admissible proof that evidence was never altered, deleted, or backdated.")
        ],
        "On our Audit Trail tab, Agent Crypt-Ledger provides the gold standard of courtroom integrity. Every action—from evidence upload to biometric verification—is recorded with an immutable block hash. This creates a tamper-proof chain of custody that is fully admissible in a court of law.",
        "Scroll through the Audit Trail timeline, filter by action (e.g. 'Biometric Verified', 'Chain of Custody Transfer'), and highlight the verified green badges."
    )

    # Slide 8: Court Reports & Live Collaboration
    add_slide(
        8,
        "Automated Forensic Dossiers & Tactical Investigation Board",
        "/reports & /collaboration (Reports & Collab Tabs)",
        "Agent Crypt-Ledger (Dossier Synthesizer)",
        [
            ("One-Click Court Dossier Generation: ", "Automatically synthesizes 3D spatial findings, deepfake heatmaps, and biometric match reports into standard legal documents."),
            ("Tactical Collaborative Board: ", "Real-time interactive canvas where forensic teams map evidence links and conduct multi-officer case reviews."),
            ("Multi-Signature Case Sign-Off: ", "Ensures formal cryptographic approval from lead investigators and lab supervisors before finalizing cases.")
        ],
        "Finally, on our Reports and Collaboration tabs, Agent Crypt-Ledger compiles findings from all modules into a court-ready forensic dossier with a single click, while our tactical board allows investigators to collaborate live and seal reports with multi-signature approvals.",
        "Show the Reports page with generated forensic summaries and click to the Collaboration tab to display the interactive team canvas."
    )

    # Slide 9: Winning Q&A Defense for Judges
    add_slide(
        9,
        "Winning Q&A Defense for Judges: Feature-Linked Answers",
        "Mastering the Panel Discussion",
        "The Complete 4-Agent Forensic Grid",
        [
            ("Q: 'How does your 3D Crime Scene work?'", "\"Agent Apex-Vision uses natural language parsing to map physical spatial coordinates in WebGL, computing inverse ballistic trajectories with millimeter precision.\""),
            ("Q: 'Why trust your Deepfake Detection?'", "\"We use a dual-model ensemble combined with Grad-CAM frequency activation heatmaps that expose sub-pixel up-sampling artifacts.\""),
            ("Q: 'How does Smart Assignment choose officers?'", "\"Agent Nexus-Decision calculates a 5-factor optimization matrix weighing specialization match, caseload capacity, distance, and historical success rate.\""),
            ("Q: 'How do you prove chain of custody in court?'", "\"Agent Crypt-Ledger records cryptographic Merkle state transitions on our audit ledger for every evidence interaction, ensuring mathematical non-repudiation.\"")
        ],
        "Judges, every advanced feature you see today is fully integrated, operational, and linked directly to our 4-agent forensic architecture. We are ready for your questions.",
        "Summarize the 4 Agents: Agent Apex-Vision, Agent Bio-Topology, Agent Nexus-Decision, and Agent Crypt-Ledger."
    )

    output_path = r"c:\Users\manoj\Downloads\orbit-oasis (1)\Forensic_AI_Executive_Presentation.docx"
    doc.save(output_path)
    print(f"Feature-linked presentation saved to: {output_path}")

if __name__ == "__main__":
    generate_presentation_docx()
